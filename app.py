"""Flask Accessibility Auditor Application.

Checks websites and documents for accessibility issues using axe-core (WCAG 2.2 AA for web pages),
PyMuPDF (for PDF documents), and python-docx (for Word documents).
"""

import json
import os
import re
import tempfile
from datetime import datetime, timezone

import fitz  # PyMuPDF
import requests
from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from flask import (
    Flask,
    flash,
    redirect,
    render_template,
    request,
    url_for,
)
from flask_sqlalchemy import SQLAlchemy
from playwright.sync_api import TimeoutError as PlaywrightTimeoutError
from playwright.sync_api import sync_playwright
from werkzeug.utils import secure_filename

app = Flask(__name__)
app.config["SECRET_KEY"] = os.environ.get("SECRET_KEY", "dev-secret-change-in-production")
app.config["SQLALCHEMY_DATABASE_URI"] = os.environ.get(
    "DATABASE_URL", "sqlite:///audits.db"
)
app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False
app.config["UPLOAD_FOLDER"] = os.path.join(tempfile.gettempdir(), "accessibility_uploads")
app.config["MAX_CONTENT_LENGTH"] = 16 * 1024 * 1024  # 16 MB max upload

ALLOWED_EXTENSIONS = {"pdf", "html", "htm", "docx"}

os.makedirs(app.config["UPLOAD_FOLDER"], exist_ok=True)

db = SQLAlchemy(app)

# Path to the bundled axe-core script served locally
AXE_CORE_PATH = os.path.join(os.path.dirname(__file__), "static", "js", "axe.min.js")


# ---------------------------------------------------------------------------
# Database Models
# ---------------------------------------------------------------------------


class Audit(db.Model):
    """Stores one accessibility audit run."""

    __tablename__ = "audits"

    id = db.Column(db.Integer, primary_key=True)
    target = db.Column(db.String(2048), nullable=False)
    audit_type = db.Column(db.String(20), nullable=False)  # "url", "pdf", "html"
    created_at = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc))
    violations_count = db.Column(db.Integer, default=0)
    passes_count = db.Column(db.Integer, default=0)
    incomplete_count = db.Column(db.Integer, default=0)
    inapplicable_count = db.Column(db.Integer, default=0)
    score = db.Column(db.Float, default=0.0)  # 0-100 accessibility score
    results_json = db.Column(db.Text, nullable=True)
    error_message = db.Column(db.Text, nullable=True)
    status = db.Column(db.String(20), default="pending")  # "success", "error"
    
    # Historical tracking fields
    parent_audit_id = db.Column(db.Integer, db.ForeignKey("audits.id"), nullable=True)
    rerun_number = db.Column(db.Integer, default=0)  # 0 = original, 1+ = rerun count
    
    # Relationships
    parent = db.relationship("Audit", remote_side=[id], backref="reruns", foreign_keys=[parent_audit_id])

    def results(self):
        """Return parsed results dict or empty dict."""
        if self.results_json:
            return json.loads(self.results_json)
        return {}
    
    def get_root_audit(self):
        """Get the original audit in this chain."""
        if self.parent_audit_id is None:
            return self
        return self.parent.get_root_audit()
    
    def get_audit_history(self):
        """Get all audits in this chain, ordered from oldest to newest."""
        root = self.get_root_audit()
        # Get all audits in the chain
        history = [root]
        history.extend(
            Audit.query.filter_by(parent_audit_id=root.id)
            .order_by(Audit.created_at.asc())
            .all()
        )
        return history

    def __repr__(self):
        return f"<Audit id={self.id} target={self.target!r} status={self.status!r}>"


# ---------------------------------------------------------------------------
# Accessibility Checking Logic
# ---------------------------------------------------------------------------


def _load_axe_script() -> str:
    """Return the axe-core JS source code."""
    with open(AXE_CORE_PATH, encoding="utf-8") as fh:
        return fh.read()


def check_url_accessibility(url: str) -> dict:
    """Run axe-core accessibility checks against a live URL using Playwright.

    Returns a dict with keys: violations, passes, incomplete, inapplicable.
    Raises RuntimeError on failure.
    """
    axe_script = _load_axe_script()

    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        try:
            context = browser.new_context(
                user_agent=(
                    "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 "
                    "(KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
                )
            )
            page = context.new_page()
            page.set_default_timeout(30_000)
            page.goto(url, wait_until="domcontentloaded")
            # Inject and run axe-core
            page.evaluate(axe_script)
            results = page.evaluate(
                """async () => {
                    const results = await axe.run(document, {
                        runOnly: { type: 'tag', values: ['wcag22aa', 'best-practice'] }
                    });
                    return {
                        violations: results.violations,
                        passes: results.passes,
                        incomplete: results.incomplete,
                        inapplicable: results.inapplicable
                    };
                }"""
            )
            return results
        except PlaywrightTimeoutError as exc:
            raise RuntimeError(f"Page timed out: {exc}") from exc
        finally:
            browser.close()


def check_html_accessibility(html_content: str, filename: str) -> dict:
    """Run axe-core against an uploaded HTML file using Playwright."""
    axe_script = _load_axe_script()

    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        try:
            page = browser.new_page()
            page.set_content(html_content, wait_until="domcontentloaded")
            page.evaluate(axe_script)
            results = page.evaluate(
                """async () => {
                    const results = await axe.run(document, {
                        runOnly: { type: 'tag', values: ['wcag22aa', 'best-practice'] }
                    });
                    return {
                        violations: results.violations,
                        passes: results.passes,
                        incomplete: results.incomplete,
                        inapplicable: results.inapplicable
                    };
                }"""
            )
            return results
        finally:
            browser.close()


def check_pdf_accessibility(pdf_path: str) -> dict:
    """Check a PDF file for common accessibility issues using PyMuPDF.

    Returns a dict with violations, passes, and metadata similar to axe-core output.
    """
    violations = []
    passes = []
    incomplete = []

    doc = fitz.open(pdf_path)
    metadata = doc.metadata or {}

    # --- Check: Document title in metadata ---
    title = metadata.get("title", "").strip()
    if title:
        passes.append(
            {
                "id": "pdf-title",
                "description": "PDF document has a title in metadata.",
                "help": "Document title",
                "helpUrl": "https://www.w3.org/TR/WCAG21/#page-titled",
                "impact": None,
                "tags": ["wcag2a", "wcag211"],
                "nodes": [],
            }
        )
    else:
        violations.append(
            {
                "id": "pdf-title",
                "impact": "serious",
                "description": "PDF document is missing a title in metadata.",
                "help": "Documents must have a title (WCAG 2.4.2).",
                "helpUrl": "https://www.w3.org/TR/WCAG21/#page-titled",
                "tags": ["wcag2a", "wcag242"],
                "nodes": [{"html": "<metadata>", "target": ["metadata"], "failureSummary": "Document title is missing from PDF metadata."}],
            }
        )

    # --- Check: Text extractability (not a scanned image) ---
    total_pages = len(doc)
    pages_with_text = sum(1 for page in doc if page.get_text("text").strip())
    if total_pages == 0:
        violations.append(
            {
                "id": "pdf-empty",
                "impact": "critical",
                "description": "PDF document appears to be empty.",
                "help": "PDF has no pages.",
                "helpUrl": "https://www.w3.org/TR/WCAG21/",
                "tags": [],
                "nodes": [],
            }
        )
    elif pages_with_text == 0:
        violations.append(
            {
                "id": "pdf-text-content",
                "impact": "critical",
                "description": (
                    "No extractable text found — PDF may consist entirely of scanned images "
                    "without OCR text layer."
                ),
                "help": "PDFs must contain real text, not just images (WCAG 1.1.1).",
                "helpUrl": "https://www.w3.org/TR/WCAG21/#non-text-content",
                "tags": ["wcag2a", "wcag111"],
                "nodes": [
                    {
                        "html": f"<pages count={total_pages}>",
                        "target": ["page"],
                        "failureSummary": "No text layer detected. Add OCR or use an accessible PDF.",
                    }
                ],
            }
        )
    else:
        text_coverage = pages_with_text / total_pages
        if text_coverage < 0.5:
            incomplete.append(
                {
                    "id": "pdf-text-content",
                    "impact": "moderate",
                    "description": (
                        f"Only {pages_with_text} of {total_pages} pages have extractable text. "
                        "Some pages may be scanned images."
                    ),
                    "help": "All pages should have accessible text.",
                    "helpUrl": "https://www.w3.org/TR/WCAG21/#non-text-content",
                    "tags": ["wcag2a", "wcag111"],
                    "nodes": [],
                }
            )
        else:
            passes.append(
                {
                    "id": "pdf-text-content",
                    "description": f"PDF has extractable text on {pages_with_text}/{total_pages} pages.",
                    "help": "Text content",
                    "helpUrl": "https://www.w3.org/TR/WCAG21/#non-text-content",
                    "impact": None,
                    "tags": ["wcag2a", "wcag111"],
                    "nodes": [],
                }
            )

    # --- Check: Tagged PDF (structure tags) ---
    is_tagged = bool(doc.get_toc()) or _pdf_has_structure_tags(doc)
    if is_tagged:
        passes.append(
            {
                "id": "pdf-tagged",
                "description": "PDF appears to have structure/tag information.",
                "help": "Tagged PDF",
                "helpUrl": "https://www.w3.org/TR/WCAG21/#info-and-relationships",
                "impact": None,
                "tags": ["wcag2a", "wcag131"],
                "nodes": [],
            }
        )
    else:
        violations.append(
            {
                "id": "pdf-tagged",
                "impact": "serious",
                "description": "PDF does not appear to be tagged (no structure tags or TOC found).",
                "help": "PDFs should be tagged for accessibility (WCAG 1.3.1).",
                "helpUrl": "https://www.w3.org/TR/WCAG21/#info-and-relationships",
                "tags": ["wcag2a", "wcag131"],
                "nodes": [
                    {
                        "html": "<pdf>",
                        "target": ["document"],
                        "failureSummary": "Add structure tags to the PDF.",
                    }
                ],
            }
        )

    # --- Check: Language metadata ---
    language = metadata.get("language", "").strip() or _detect_pdf_language_marker(doc)
    if language:
        passes.append(
            {
                "id": "pdf-language",
                "description": f"PDF document language is set to: {language}",
                "help": "Document language",
                "helpUrl": "https://www.w3.org/TR/WCAG21/#language-of-page",
                "impact": None,
                "tags": ["wcag2a", "wcag311"],
                "nodes": [],
            }
        )
    else:
        violations.append(
            {
                "id": "pdf-language",
                "impact": "serious",
                "description": "PDF document language is not specified in metadata.",
                "help": "The language of the document must be specified (WCAG 3.1.1).",
                "helpUrl": "https://www.w3.org/TR/WCAG21/#language-of-page",
                "tags": ["wcag2a", "wcag311"],
                "nodes": [
                    {
                        "html": "<metadata>",
                        "target": ["metadata"],
                        "failureSummary": "Set the document language in PDF properties.",
                    }
                ],
            }
        )

    # --- Check: Hyperlinks have descriptive text ---
    links_checked, bad_links = _check_pdf_links(doc)
    if links_checked > 0:
        if bad_links:
            violations.append(
                {
                    "id": "pdf-link-text",
                    "impact": "moderate",
                    "description": f"{len(bad_links)} link(s) have missing or non-descriptive text.",
                    "help": "Link text must be descriptive (WCAG 2.4.4).",
                    "helpUrl": "https://www.w3.org/TR/WCAG21/#link-purpose-in-context",
                    "tags": ["wcag2a", "wcag244"],
                    "nodes": [
                        {
                            "html": f"<a href={url!r}>",
                            "target": ["a"],
                            "failureSummary": "Link is missing descriptive text.",
                        }
                        for url in bad_links[:10]  # cap at 10 for display
                    ],
                }
            )
        else:
            passes.append(
                {
                    "id": "pdf-link-text",
                    "description": f"All {links_checked} link(s) appear to have associated text.",
                    "help": "Link text",
                    "helpUrl": "https://www.w3.org/TR/WCAG21/#link-purpose-in-context",
                    "impact": None,
                    "tags": ["wcag2a", "wcag244"],
                    "nodes": [],
                }
            )

    doc.close()
    return {
        "violations": violations,
        "passes": passes,
        "incomplete": incomplete,
        "inapplicable": [],
    }


def check_docx_accessibility(docx_path: str) -> dict:
    """Check a Word (.docx) document for common accessibility issues.
    
    Returns a dict with violations, passes, and metadata similar to axe-core output.
    """
    violations = []
    passes = []
    incomplete = []
    
    doc = DocxDocument(docx_path)
    core_props = doc.core_properties
    
    # --- Check: Document title ---
    title = core_props.title
    if title and title.strip():
        passes.append({
            "id": "docx-title",
            "description": "Word document has a title in properties.",
            "help": "Document title",
            "helpUrl": "https://www.w3.org/TR/WCAG22/#page-titled",
            "impact": None,
            "tags": ["wcag2a", "wcag242"],
            "nodes": [],
        })
    else:
        violations.append({
            "id": "docx-title",
            "impact": "serious",
            "description": "Word document is missing a title in document properties.",
            "help": "Documents must have a title (WCAG 2.4.2).",
            "helpUrl": "https://www.w3.org/TR/WCAG22/#page-titled",
            "tags": ["wcag2a", "wcag242"],
            "nodes": [{
                "html": "<properties>",
                "target": ["properties"],
                "failureSummary": "Document title is missing from Word properties.",
            }],
        })
    
    # --- Check: Document language ---
    language = core_props.language
    if language and language.strip():
        passes.append({
            "id": "docx-language",
            "description": f"Word document language is set to: {language}",
            "help": "Document language",
            "helpUrl": "https://www.w3.org/TR/WCAG22/#language-of-page",
            "impact": None,
            "tags": ["wcag2a", "wcag311"],
            "nodes": [],
        })
    else:
        violations.append({
            "id": "docx-language",
            "impact": "serious",
            "description": "Word document language is not specified.",
            "help": "The language of the document must be specified (WCAG 3.1.1).",
            "helpUrl": "https://www.w3.org/TR/WCAG22/#language-of-page",
            "tags": ["wcag2a", "wcag311"],
            "nodes": [{
                "html": "<properties>",
                "target": ["properties"],
                "failureSummary": "Set the document language in Word properties.",
            }],
        })
    
    # --- Check: Headings structure ---
    headings = []
    for para in doc.paragraphs:
        if para.style.name.startswith('Heading'):
            headings.append(para.style.name)
    
    if headings:
        passes.append({
            "id": "docx-headings",
            "description": f"Document uses heading styles ({len(headings)} heading(s) found).",
            "help": "Document structure",
            "helpUrl": "https://www.w3.org/TR/WCAG22/#info-and-relationships",
            "impact": None,
            "tags": ["wcag2a", "wcag131"],
            "nodes": [],
        })
    else:
        # Check if document has any text content
        has_content = any(para.text.strip() for para in doc.paragraphs)
        if has_content:
            violations.append({
                "id": "docx-headings",
                "impact": "moderate",
                "description": "No heading styles detected. Use heading styles for document structure.",
                "help": "Document should use heading styles for organization (WCAG 1.3.1).",
                "helpUrl": "https://www.w3.org/TR/WCAG22/#info-and-relationships",
                "tags": ["wcag2a", "wcag131"],
                "nodes": [{
                    "html": "<document>",
                    "target": ["document"],
                    "failureSummary": "Apply heading styles to organize content.",
                }],
            })
    
    # --- Check: Images have alt text ---
    images_checked = 0
    images_without_alt = 0
    
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            images_checked += 1
    
    # Count inline shapes (images with descriptions)
    for section in doc.sections:
        for para in doc.paragraphs:
            if para._element.xpath('.//pic:pic', namespaces={
                'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture'
            }):
                # This is a simplified check - proper implementation would check alt text
                # For now, we'll flag this for manual review
                pass
    
    if images_checked > 0:
        incomplete.append({
            "id": "docx-images",
            "impact": "serious",
            "description": f"{images_checked} image(s) found. Manual review required to verify alt text.",
            "help": "All images must have alternative text (WCAG 1.1.1).",
            "helpUrl": "https://www.w3.org/TR/WCAG22/#non-text-content",
            "tags": ["wcag2a", "wcag111"],
            "nodes": [],
        })
    else:
        passes.append({
            "id": "docx-images",
            "description": "No images detected in document.",
            "help": "Non-text content",
            "helpUrl": "https://www.w3.org/TR/WCAG22/#non-text-content",
            "impact": None,
            "tags": ["wcag2a", "wcag111"],
            "nodes": [],
        })
    
    # --- Check: Hyperlinks ---
    hyperlinks = []
    for para in doc.paragraphs:
        if para._element.xpath('.//w:hyperlink', namespaces={
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        }):
            hyperlinks.append(para)
    
    if hyperlinks:
        # Check if links have descriptive text (manual review needed)
        incomplete.append({
            "id": "docx-links",
            "impact": "moderate",
            "description": f"{len(hyperlinks)} hyperlink(s) found. Verify they have descriptive text.",
            "help": "Link text must be descriptive (WCAG 2.4.4).",
            "helpUrl": "https://www.w3.org/TR/WCAG22/#link-purpose-in-context",
            "tags": ["wcag2a", "wcag244"],
            "nodes": [],
        })
    
    # --- Check: Tables ---
    tables = doc.tables
    if tables:
        incomplete.append({
            "id": "docx-tables",
            "impact": "moderate",
            "description": f"{len(tables)} table(s) found. Verify they have header rows defined.",
            "help": "Tables must have proper header structure (WCAG 1.3.1).",
            "helpUrl": "https://www.w3.org/TR/WCAG22/#info-and-relationships",
            "tags": ["wcag2a", "wcag131"],
            "nodes": [],
        })
    
    # --- Check: Document has content ---
    total_text = "".join(para.text for para in doc.paragraphs).strip()
    if not total_text:
        violations.append({
            "id": "docx-content",
            "impact": "critical",
            "description": "Document appears to be empty or has no text content.",
            "help": "Document must have content.",
            "helpUrl": "https://www.w3.org/TR/WCAG22/",
            "tags": [],
            "nodes": [],
        })
    else:
        passes.append({
            "id": "docx-content",
            "description": "Document contains text content.",
            "help": "Document content",
            "helpUrl": "https://www.w3.org/TR/WCAG22/",
            "impact": None,
            "tags": [],
            "nodes": [],
        })
    
    return {
        "violations": violations,
        "passes": passes,
        "incomplete": incomplete,
        "inapplicable": [],
    }


def _pdf_has_structure_tags(doc: fitz.Document) -> bool:
    """Return True if the PDF has structure/accessibility tags."""
    try:
        xml = doc.get_xml_metadata()
        if xml and ("pdfaSchema" in xml or "Marked" in xml):
            return True
    except Exception:
        pass
    return False


def _detect_pdf_language_marker(doc: fitz.Document) -> str:
    """Try to detect language from XML metadata."""
    try:
        xml = doc.get_xml_metadata()
        if xml:
            match = re.search(r'xml:lang=["\']([^"\']+)["\']', xml)
            if match:
                return match.group(1)
            match = re.search(r'<dc:language[^>]*>([^<]+)</dc:language>', xml)
            if match:
                return match.group(1)
    except Exception:
        pass
    return ""


def _check_pdf_links(doc: fitz.Document):
    """Return (total_links, list_of_bare_urls_without_text)."""
    bare_urls = []
    total = 0
    for page in doc:
        for link in page.get_links():
            if link.get("kind") == fitz.LINK_URI:
                total += 1
                uri = link.get("uri", "")
                # Check if the link rectangle has visible text on the page
                rect = fitz.Rect(link.get("from"))
                words = page.get_text("words", clip=rect)
                if not words:
                    bare_urls.append(uri)
    return total, bare_urls


# ---------------------------------------------------------------------------
# Score Calculation
# ---------------------------------------------------------------------------


def _compute_score(violations: list, passes: list, incomplete: list) -> float:
    """Compute a 0-100 accessibility score.

    Higher is better. Weighted by impact level.
    """
    impact_weights = {"critical": 4, "serious": 3, "moderate": 2, "minor": 1}

    penalty = sum(
        impact_weights.get(v.get("impact", "minor"), 1) for v in violations
    )
    bonus = len(passes)
    total_signals = penalty + bonus + len(incomplete)
    if total_signals == 0:
        return 100.0
    raw = bonus / (bonus + penalty + 0.5 * len(incomplete))
    return round(raw * 100, 1)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def allowed_file(filename: str) -> bool:
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


def _normalise_url(url: str) -> str:
    url = url.strip()
    if not url.startswith(("http://", "https://")):
        url = "https://" + url
    return url


def _save_audit(target: str, audit_type: str, results: dict | None, error: str | None, 
                parent_audit_id: int | None = None, rerun_number: int = 0) -> Audit:
    """Persist an audit to the database and return the model instance."""
    audit = Audit(
        target=target,
        audit_type=audit_type,
        parent_audit_id=parent_audit_id,
        rerun_number=rerun_number,
    )
    if error:
        audit.status = "error"
        audit.error_message = error
    else:
        violations = results.get("violations", [])
        passes = results.get("passes", [])
        incomplete = results.get("incomplete", [])
        inapplicable = results.get("inapplicable", [])
        audit.status = "success"
        audit.violations_count = len(violations)
        audit.passes_count = len(passes)
        audit.incomplete_count = len(incomplete)
        audit.inapplicable_count = len(inapplicable)
        audit.score = _compute_score(violations, passes, incomplete)
        audit.results_json = json.dumps(results)

    db.session.add(audit)
    db.session.commit()
    return audit


def convert_to_markdown(url: str) -> dict:
    """Convert a URL to markdown using Jina AI Reader API.
    
    Returns a dict with 'markdown' content and 'error' if any.
    """
    try:
        # Jina AI Reader API endpoint
        jina_url = f"https://r.jina.ai/{url}"
        
        # Add headers as recommended by Jina
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36',
            'Accept': 'text/plain, */*'
        }
        
        response = requests.get(jina_url, headers=headers, timeout=60)
        response.raise_for_status()
        
        # Check if we got content
        if not response.text or len(response.text.strip()) == 0:
            return {
                "markdown": None,
                "error": "Jina AI Reader returned empty content. The URL may be inaccessible or blocked."
            }
        
        return {
            "markdown": response.text,
            "error": None
        }
    except requests.Timeout:
        return {
            "markdown": None,
            "error": "Request timed out. The page may be too large or slow to respond."
        }
    except requests.HTTPError as exc:
        return {
            "markdown": None,
            "error": f"HTTP error {exc.response.status_code}: Unable to convert URL."
        }
    except requests.RequestException as exc:
        return {
            "markdown": None,
            "error": f"Failed to convert URL to markdown: {str(exc)}"
        }
    except Exception as exc:
        return {
            "markdown": None,
            "error": f"Unexpected error during conversion: {str(exc)}"
        }


# ---------------------------------------------------------------------------
# Routes
# ---------------------------------------------------------------------------


@app.route("/", methods=["GET"])
def index():
    audits = Audit.query.order_by(Audit.created_at.desc()).limit(50).all()
    return render_template("index.html", audits=audits)


@app.route("/health", methods=["GET"])
def health_check():
    """Health check endpoint for Docker and monitoring."""
    return {"status": "healthy", "service": "accessibility-auditor"}, 200


@app.route("/audit/url", methods=["POST"])
def audit_url():
    url = request.form.get("url", "").strip()
    if not url:
        flash("Please enter a URL.", "warning")
        return redirect(url_for("index"))

    url = _normalise_url(url)

    try:
        results = check_url_accessibility(url)
        audit = _save_audit(url, "url", results, None)
    except (RuntimeError, PlaywrightTimeoutError, OSError, ValueError) as exc:
        audit = _save_audit(url, "url", None, str(exc))
        flash(f"Error checking {url}: {exc}", "danger")
        return redirect(url_for("index"))

    return redirect(url_for("audit_detail", audit_id=audit.id))


@app.route("/audit/file", methods=["POST"])
def audit_file():
    if "file" not in request.files:
        flash("No file part in the request.", "warning")
        return redirect(url_for("index"))

    uploaded_file = request.files["file"]
    if not uploaded_file or uploaded_file.filename == "":
        flash("No file selected.", "warning")
        return redirect(url_for("index"))

    if not allowed_file(uploaded_file.filename):
        flash("Unsupported file type. Please upload a PDF, HTML, or Word document.", "warning")
        return redirect(url_for("index"))

    filename = secure_filename(uploaded_file.filename)
    ext = filename.rsplit(".", 1)[1].lower()
    if ext == "pdf":
        audit_type = "pdf"
    elif ext == "docx":
        audit_type = "docx"
    else:  # html / htm
        audit_type = "html"
    tmp_path = os.path.join(app.config["UPLOAD_FOLDER"], filename)
    uploaded_file.save(tmp_path)

    try:
        if ext == "pdf":
            results = check_pdf_accessibility(tmp_path)
        elif ext == "docx":
            results = check_docx_accessibility(tmp_path)
        else:  # html / htm
            with open(tmp_path, encoding="utf-8", errors="replace") as fh:
                html_content = fh.read()
            results = check_html_accessibility(html_content, filename)

        audit = _save_audit(filename, audit_type, results, None)
    except (RuntimeError, PlaywrightTimeoutError, OSError, ValueError, fitz.FileDataError, Exception) as exc:
        audit = _save_audit(filename, audit_type, None, str(exc))
        flash(f"Error processing {filename}: {exc}", "danger")
        return redirect(url_for("index"))
    finally:
        if os.path.exists(tmp_path):
            os.remove(tmp_path)

    return redirect(url_for("audit_detail", audit_id=audit.id))


@app.route("/audit/<int:audit_id>")
def audit_detail(audit_id: int):
    audit = db.get_or_404(Audit, audit_id)
    results = audit.results()
    history = audit.get_audit_history()
    return render_template("audit_detail.html", audit=audit, results=results, history=history)


@app.route("/audit/<int:audit_id>/rerun", methods=["POST"])
def rerun_audit(audit_id: int):
    """Rerun an audit with the same target and type, linking to the original."""
    original_audit = db.get_or_404(Audit, audit_id)
    
    # Get the root audit to maintain a single chain
    root_audit = original_audit.get_root_audit()
    
    # Calculate the next rerun number
    existing_reruns = Audit.query.filter_by(parent_audit_id=root_audit.id).count()
    next_rerun_number = existing_reruns + 1
    
    # Rerun based on audit type
    try:
        if original_audit.audit_type == "url":
            url = _normalise_url(original_audit.target)
            results = check_url_accessibility(url)
            new_audit = _save_audit(
                url, 
                "url", 
                results, 
                None,
                parent_audit_id=root_audit.id,
                rerun_number=next_rerun_number
            )
        elif original_audit.audit_type == "pdf":
            flash("PDF re-auditing requires re-uploading the file.", "warning")
            return redirect(url_for("audit_detail", audit_id=audit_id))
        elif original_audit.audit_type == "docx":
            flash("Word document re-auditing requires re-uploading the file.", "warning")
            return redirect(url_for("audit_detail", audit_id=audit_id))
        elif original_audit.audit_type == "html":
            flash("HTML file re-auditing requires re-uploading the file.", "warning")
            return redirect(url_for("audit_detail", audit_id=audit_id))
        else:
            flash(f"Cannot rerun audit of type {original_audit.audit_type}.", "danger")
            return redirect(url_for("audit_detail", audit_id=audit_id))
        
        flash(f"Audit rerun completed successfully! (Run #{next_rerun_number})", "success")
        return redirect(url_for("audit_detail", audit_id=new_audit.id))
    
    except (RuntimeError, PlaywrightTimeoutError, OSError, ValueError) as exc:
        new_audit = _save_audit(
            original_audit.target,
            original_audit.audit_type,
            None,
            str(exc),
            parent_audit_id=root_audit.id,
            rerun_number=next_rerun_number
        )
        flash(f"Error during rerun: {exc}", "danger")
        return redirect(url_for("audit_detail", audit_id=new_audit.id))


@app.route("/convert/url", methods=["POST"])
def convert_url_to_markdown():
    """Convert a URL to markdown using Jina AI Reader."""
    url = request.form.get("url", "").strip()
    
    if not url:
        flash("Please enter a URL to convert.", "warning")
        return redirect(url_for("index"))
    
    url = _normalise_url(url)
    result = convert_to_markdown(url)
    
    if result["error"]:
        flash(result["error"], "danger")
        return redirect(url_for("index"))
    
    return render_template("markdown_viewer.html", url=url, markdown=result["markdown"])


@app.route("/audit/<int:audit_id>/convert-markdown", methods=["POST"])
def convert_audit_to_markdown(audit_id: int):
    """Convert an audited URL to markdown."""
    audit = db.get_or_404(Audit, audit_id)
    
    if audit.audit_type != "url":
        flash("Only URL audits can be converted to markdown.", "warning")
        return redirect(url_for("audit_detail", audit_id=audit_id))
    
    result = convert_to_markdown(audit.target)
    
    if result["error"]:
        flash(result["error"], "danger")
        return redirect(url_for("audit_detail", audit_id=audit_id))
    
    return render_template("markdown_viewer.html", url=audit.target, markdown=result["markdown"], audit_id=audit_id)


@app.route("/audit/<int:audit_id>/delete", methods=["POST"])
def delete_audit(audit_id: int):
    audit = db.get_or_404(Audit, audit_id)
    db.session.delete(audit)
    db.session.commit()
    flash("Audit deleted.", "info")
    return redirect(url_for("index"))


# ---------------------------------------------------------------------------
# App Entry Point
# ---------------------------------------------------------------------------


with app.app_context():
    db.create_all()

if __name__ == "__main__":
    debug_mode = os.environ.get("FLASK_DEBUG", "0") == "1"
    app.run(debug=debug_mode, host="0.0.0.0", port=5000)
